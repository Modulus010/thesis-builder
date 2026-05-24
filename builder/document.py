#!/usr/bin/env python3

import io
import os
import re
import subprocess
import sys
from typing import Dict
try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from ast_nodes import Figure, Table, PlantUMLBlock, CodeBlock, Section, Thesis, CITATION_DISPLAY_RE
from builder.styles import StyleConfig
from builder.xml_helpers import (
    ensure_trPr, set_row_cant_split, set_row_table_header, set_spacing_lines,
    make_border, find_drawing_and_inline, build_anchor_from_inline,
    remove_theme_fonts, remove_child_tags,
)
from builder.numbering import collect_numbering, caption_with_number


CJK_FONTS = {"宋体", "黑体", "楷体", "仿宋"}


class DocumentBuilder:

    _T_DECLARATION = "郑 重 声 明"
    _T_ABSTRACT_CN = "摘  要"
    _T_ABSTRACT_EN = "ABSTRACT"
    _T_TOC = "目  录"
    _T_REFERENCES = "参 考 文 献"
    _T_APPENDIX = "附  录"
    _T_ACKNOWLEDGMENTS = "致  谢"

    def __init__(self, config=None):
        self.styles = config if isinstance(config, StyleConfig) else StyleConfig(config if config else {})
        self.doc = None
        self.thesis = None
        self.figures_dir = ""
        self._body_font = self.styles.font("body")
        page = self.styles.page
        self._content_width_cm = page.get("width", 21.0) - page.get("margin_left", 3.0) - page.get("margin_right", 2.5)
        self._content_height_cm = page.get("height", 29.7) - page.get("margin_top", 2.5) - page.get("margin_bottom", 2.5)

    def _is_body_default(self, font_cfg: Dict) -> bool:
        return (font_cfg.get("name") == self._body_font.get("name")
                and font_cfg.get("size") == self._body_font.get("size")
                and font_cfg.get("bold") == self._body_font.get("bold"))

    def _is_body_font_family(self, font_cfg: Dict) -> bool:
        return font_cfg.get("name") == self._body_font.get("name")

    @staticmethod
    def _set_rfonts_on_rpr(rPr, east_asia: str, ascii_font: str = None, h_ansi: str = None):
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), east_asia)
        if ascii_font:
            rFonts.set(qn("w:ascii"), ascii_font)
        if h_ansi:
            rFonts.set(qn("w:hAnsi"), h_ansi)

    def _set_rfonts(self, run, east_asia: str, ascii_font: str = None, h_ansi: str = None):
        self._set_rfonts_on_rpr(run._element.get_or_add_rPr(), east_asia, ascii_font, h_ansi)

    def _apply_run_font(self, run, font_cfg: Dict):
        font_name = font_cfg.get("name", "宋体")
        if self._is_body_default(font_cfg):
            return
        bold = font_cfg.get("bold", False)
        if bold:
            run.font.bold = True
        if self._is_body_font_family(font_cfg):
            size = font_cfg.get("size", 12)
            if size != self._body_font.get("size", 12):
                run.font.size = Pt(size)
        else:
            run.font.name = font_name
            run.font.size = Pt(font_cfg.get("size", 12))
            ea = font_name if font_name in CJK_FONTS else "宋体"
            self._set_rfonts(run, east_asia=ea, ascii_font="Times New Roman",
                             h_ansi="Times New Roman")

    def _set_paragraph_text(self, para, text: str, font_key: str = "body",
                            with_line_spacing: bool = True):
        run = para.add_run(text)
        self._apply_run_font(run, self.styles.font(font_key))
        if with_line_spacing:
            self._apply_line_spacing(para)
        return para

    def _apply_line_spacing(self, para):
        para.paragraph_format.line_spacing = Pt(self.styles.layout.line_spacing_pt)

    def _add_body_with_citations(self, para, text: str):
        body_font = self.styles.font("body")
        parts = CITATION_DISPLAY_RE.split(text)
        for part in parts:
            if not part:
                continue
            if CITATION_DISPLAY_RE.fullmatch(part):
                self._add_citation_hyperlink(para, part, body_font)
            else:
                run = para.add_run(part)
                self._apply_run_font(run, body_font)

    def _add_citation_hyperlink(self, para, citation: str, font_cfg: dict):
        inner = citation[1:-1]
        self._add_superscript_run(para, '[', font_cfg)
        i = 0
        while i < len(inner):
            ch = inner[i]
            if ch in (',', '，'):
                self._add_superscript_run(para, ch, font_cfg)
                i += 1
            elif ch.isdigit():
                j = i
                while j < len(inner) and inner[j].isdigit():
                    j += 1
                num = inner[i:j]
                self._add_ref_hyperlink(para, num, font_cfg)
                i = j
                if i < len(inner) and inner[i] == '-':
                    self._add_superscript_run(para, '-', font_cfg)
                    i += 1
            else:
                self._add_superscript_run(para, ch, font_cfg)
                i += 1
        self._add_superscript_run(para, ']', font_cfg)

    def _add_superscript_run(self, para, text: str, font_cfg: dict):
        run = para.add_run(text)
        self._apply_run_font(run, font_cfg)
        run.font.superscript = True

    def _add_ref_hyperlink(self, para, ref_num: str, font_cfg: dict):
        hl = OxmlElement('w:hyperlink')
        hl.set(qn('w:anchor'), f'_Ref{ref_num}')
        hl.set(qn('w:history'), '1')

        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)

        if not self._is_body_default(font_cfg):
            fn = font_cfg.get('name', '宋体')
            size_val = str(int(font_cfg.get('size', 12) * 2))
            rFonts = OxmlElement('w:rFonts')
            if fn in CJK_FONTS:
                rFonts.set(qn('w:eastAsia'), fn)
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), size_val)
            rPr.append(sz)

        vertAlign = OxmlElement('w:vertAlign')
        vertAlign.set(qn('w:val'), 'superscript')
        rPr.append(vertAlign)

        r.append(rPr)

        t = OxmlElement('w:t')
        t.text = ref_num
        r.append(t)

        hl.append(r)
        para._element.append(hl)

    def _add_section(self, start_type: int):
        section = self.doc.add_section(start_type)
        self._setup_section_margins(section)
        return section

    @staticmethod
    def _cm_to_twips(cm_val: float) -> int:
        return round(cm_val / 2.54 * 1440)

    def _setup_section_margins(self, section):
        page_cfg = self.styles.page
        c2t = self._cm_to_twips
        top = c2t(float(page_cfg.get("margin_top", 2.5)))
        bottom = c2t(float(page_cfg.get("margin_bottom", 2.5)))
        left = c2t(float(page_cfg.get("margin_left", 3.0)))
        right = c2t(float(page_cfg.get("margin_right", 2.5)))
        header = c2t(float(page_cfg.get("header_distance", 1.5)))
        footer = c2t(float(page_cfg.get("footer_distance", 1.5)))
        pgMar = section._sectPr.find(qn('w:pgMar'))
        if pgMar is None:
            pgMar = OxmlElement('w:pgMar')
            section._sectPr.insert(0, pgMar)
        pgMar.set(qn('w:top'), str(top))
        pgMar.set(qn('w:bottom'), str(bottom))
        pgMar.set(qn('w:left'), str(left))
        pgMar.set(qn('w:right'), str(right))
        pgMar.set(qn('w:header'), str(header))
        pgMar.set(qn('w:footer'), str(footer))
        pgMar.set(qn('w:gutter'), '0')
        section.page_width = Cm(float(page_cfg.get("width", 21.0)))
        section.page_height = Cm(float(page_cfg.get("height", 29.7)))
        sectPr = section._sectPr
        docGrid = sectPr.find(qn("w:docGrid"))
        if docGrid is not None:
            sectPr.remove(docGrid)

    def _text_width_twips(self) -> int:
        return self._cm_to_twips(self._content_width_cm)

    def _setup_header(self, section, chapter_title: str = None):
        section.header.is_linked_to_previous = False
        para = section.header.paragraphs[0]
        para.text = ""

        # Remove Header style to avoid its default tab stops interfering
        pPr = para._element.get_or_add_pPr()
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is not None:
            pPr.remove(pStyle)

        left_run = para.add_run("东北大学本科生毕业设计（论文）")
        self._apply_run_font(left_run, self.styles.font("header_footer"))

        if chapter_title:
            para.add_run("\t")
            right_run = para.add_run(chapter_title)
            self._apply_run_font(right_run, self.styles.font("header_footer"))

        self._add_underline(para)

        if chapter_title:
            self._add_right_tab_stop(para)

    def _setup_footer(self, section):
        section.footer.is_linked_to_previous = False
        para = section.footer.paragraphs[0]
        para.text = ""
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._add_page_number_field(para, self.styles.font("page_number"))

    def _set_page_number_format(self, section, fmt: str, start: int = None):
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn("w:pgNumType"))
        if pgNumType is None:
            pgNumType = OxmlElement("w:pgNumType")
            sectPr.append(pgNumType)
        pgNumType.set(qn("w:fmt"), fmt)
        if start is not None:
            pgNumType.set(qn("w:start"), str(start))
        else:
            if pgNumType.get(qn("w:start")) is not None:
                del pgNumType.attrib[qn("w:start")]

    def _add_underline(self, para):
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_right_tab_stop(self, para):
        pPr = para._element.get_or_add_pPr()
        tabs = pPr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            pPr.append(tabs)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(self._text_width_twips()))
        tabs.append(tab)

    def _add_title_paragraph(self, text: str, font_key: str,
                             style: str = "Heading 1",
                             alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
                             space_before: int = None,
                             space_after: int = None):
        para = self.doc.add_paragraph(text, style=style)
        para.alignment = alignment
        self._apply_line_spacing(para)
        if style not in self._HEADING_STYLES:
            for run in para.runs:
                self._apply_run_font(run, self.styles.font(font_key))
        set_spacing_lines(
            para,
            space_before if space_before is not None else self.styles.layout.chapter_before_lines,
            space_after if space_after is not None else self.styles.layout.chapter_after_lines,
        )
        return para

    def _add_page_number_field(self, para, font_cfg: Dict):
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run = para.add_run()
        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_end)
        self._apply_run_font(run, font_cfg)


    def _apply_style_font(self, style, font_key: str, alignment=None):
        fc = self.styles.font(font_key)
        fn = fc.get("name", "宋体")
        style.font.name = fn
        style.font.size = Pt(fc.get("size", 12))
        bold = fc.get("bold", False)
        if bold:
            style.font.bold = True
        rPr = style.element.get_or_add_rPr()
        if not bold:
            remove_child_tags(rPr, qn("w:b"), qn("w:bCs"))
        self._set_rfonts_on_rpr(rPr, fn, "Times New Roman", "Times New Roman")
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            remove_theme_fonts(rFonts)
        remove_child_tags(rPr, qn("w:color"))
        if alignment is not None:
            style.paragraph_format.alignment = alignment

    def _define_heading_styles(self):
        configs = [
            ("Heading 1", "chapter_title", WD_PARAGRAPH_ALIGNMENT.CENTER),
            ("Heading 2", "section_title_1", WD_PARAGRAPH_ALIGNMENT.JUSTIFY),
            ("Heading 3", "section_title_2", WD_PARAGRAPH_ALIGNMENT.JUSTIFY),
        ]
        for style_name, font_key, alignment in configs:
            self._apply_style_font(self.doc.styles[style_name], font_key, alignment)

    def _define_caption_style(self):
        style = self.doc.styles["Caption"]
        self._apply_style_font(style, "figure_caption", WD_PARAGRAPH_ALIGNMENT.CENTER)
        fmt = style.paragraph_format
        fmt.line_spacing = Pt(self.styles.layout.line_spacing_pt)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)

    def _define_toc_styles(self):
        toc_configs = [
            (1, "toc_chapter"),
            (2, "toc_other"),
            (3, "toc_other"),
        ]
        toc_indent_pt = {2: 21, 3: 42}
        for level, font_key in toc_configs:
            style_name = f"TOC {level}"
            try:
                style = self.doc.styles[style_name]
            except KeyError:
                style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            self._apply_style_font(style, font_key)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            indent = toc_indent_pt.get(level, 0)
            if indent:
                style.paragraph_format.left_indent = Pt(indent)
                style.paragraph_format.first_line_indent = Pt(0)

    def _clean_all_theme_fonts(self):
        for style in self.doc.styles:
            rPr = style.element.find(qn("w:rPr"))
            if rPr is not None:
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is not None:
                    remove_theme_fonts(rFonts)
                remove_child_tags(rPr, qn("w:szCs"), qn("w:bCs"))

    def _define_normal_style(self):
        self._apply_style_font(self.doc.styles["Normal"], "body")

    def _strip_empty_pPr_rPr(self):
        for p in self.doc.paragraphs:
            pPr = p._element.find(qn("w:pPr"))
            if pPr is not None:
                rPr = pPr.find(qn("w:rPr"))
                if rPr is not None and len(rPr) == 0:
                    pPr.remove(rPr)

    def _add_document_section(self, start_type, header_title: str,
                               page_format: str = "decimal", page_start: int = None):
        sec = self._add_section(start_type)
        self._setup_header(sec, header_title)
        self._setup_footer(sec)
        self._set_page_number_format(sec, page_format, page_start)
        return sec

    def build(self, thesis: Thesis, output_path: str):
        self.doc = Document()
        self.thesis = thesis

        collect_numbering(thesis)

        if thesis.basedir:
            self.figures_dir = os.path.join(thesis.basedir, self.styles.figures_path)

        self._define_heading_styles()
        self._define_toc_styles()
        if self.styles.use_native_caption:
            self._define_caption_style()
        self._clean_all_theme_fonts()
        self._define_normal_style()

        self._setup_section_margins(self.doc.sections[0])

        self._build_cover()

        self._add_section(WD_SECTION_START.ODD_PAGE)
        self._build_formal_cover()

        self._add_section(WD_SECTION_START.ODD_PAGE)
        self._build_english_cover()

        sec_decl = self._add_section(WD_SECTION_START.ODD_PAGE)
        self._setup_header(sec_decl, self._T_DECLARATION)
        self._build_declaration()

        self._add_document_section(WD_SECTION_START.ODD_PAGE, self._T_ABSTRACT_CN,
                                   "upperRoman", 1)
        self._build_abstract_section(
            title=self._T_ABSTRACT_CN, title_font="abstract_title",
            lines=self.thesis.abstract, body_font="abstract_body",
            indent=True, justify=True,
            kw_label="关键词：", kw_label_font="keywords_title",
            keywords=self.thesis.abstract_keywords, kw_font="keywords",
            kw_sep="；",
        )

        self._add_document_section(WD_SECTION_START.ODD_PAGE, self._T_ABSTRACT_EN,
                                   "upperRoman")
        self._build_abstract_section(
            title=self._T_ABSTRACT_EN, title_font="english_abstract_title",
            lines=self.thesis.english_abstract, body_font="english_abstract_body",
            indent=True, justify=True,
            kw_label="Key words: ", kw_label_font="english_keywords_title",
            keywords=self.thesis.english_keywords, kw_font="english_keywords",
            kw_sep="; ",
        )

        self._add_document_section(WD_SECTION_START.ODD_PAGE, self._T_TOC,
                                   "upperRoman")
        self._build_toc()

        if self.thesis.sections:
            for i, section in enumerate(self.thesis.sections):
                header_text = f"{section.auto_number} {section.title}" if section.auto_number and section.level == 1 else section.title
                self._add_document_section(WD_SECTION_START.ODD_PAGE, header_text, "decimal",
                                           page_start=1 if i == 0 else None)
                self._build_section(section)

        self._add_document_section(WD_SECTION_START.ODD_PAGE, self._T_REFERENCES)
        self._build_references()

        if self.thesis.appendix_sections:
            self._add_document_section(WD_SECTION_START.ODD_PAGE, self._T_APPENDIX)
            self._build_appendix()

        self._add_document_section(WD_SECTION_START.ODD_PAGE, self._T_ACKNOWLEDGMENTS)
        self._build_acknowledgments()

        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        self._strip_empty_pPr_rPr()
        self.doc.save(output_path)


    def _add_full_page_image(self, image_name: str, *, behind_doc: bool = True,
                              allow_overlap: bool = True, center: bool = True,
                              include_effect_extent: bool = True):
        img_path = os.path.join(os.path.dirname(__file__), "..", "figures", image_name)
        img_path = os.path.normpath(img_path)

        page_w = Cm(self.styles.page.get("width", 21))
        page_h = Cm(self.styles.page.get("height", 29.7))

        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        if center:
            para.paragraph_format.line_spacing = Pt(1)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = para.add_run()
        run.font.size = Pt(1)
        try:
            run.add_picture(img_path, width=page_w, height=page_h)
        except OSError:
            print(f"Warning: back cover image not found: {img_path}", file=sys.stderr)
            return

        drawing_elem, inline = find_drawing_and_inline(run._element)
        if drawing_elem is None or inline is None:
            return

        anchor = build_anchor_from_inline(
            inline, behind_doc=behind_doc, allow_overlap=allow_overlap,
            include_effect_extent=include_effect_extent,
        )
        drawing_elem.replace(inline, anchor)

    def _add_cover_page(self, image_name: str):
        self._add_full_page_image(image_name, behind_doc=True, allow_overlap=True)

    def _build_cover(self):
        self._add_cover_page("cover_image1.jpeg")

        meta = self.thesis.metadata

        self._add_cover_spacer(210)

        self._add_cover_detail_line("论文题目")
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_para.paragraph_format.line_spacing = Pt(self.styles.layout.cover_line_spacing_pt)
        title_run = title_para.add_run(meta.title)
        self._apply_run_font(title_run, self.styles.font("cover_thesis_title"))

        self._add_cover_spacer(24)

        self._add_cover_info_table([
            ("学院名称", meta.college),
            ("专业名称", meta.major),
            ("学生姓名", meta.student_name),
            ("指导教师", meta.advisor),
        ])

        self._add_cover_spacer(24)

        cn_year = _chinese_year(meta.year)
        cn_month = _chinese_month(meta.month)
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date_run = date_para.add_run(f"{cn_year}年 {cn_month}月")
        self._apply_run_font(date_run, self.styles.font("cover_date"))

    def _build_formal_cover(self):
        meta = self.thesis.metadata

        id_para = self.doc.add_paragraph()
        id_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        id_run = id_para.add_run(f"学号  {meta.student_id}                密级  ")
        self._apply_run_font(id_run, self.styles.font("cover_id"))

        self._add_cover_spacer(self.styles.layout.cover_spacer_id)

        uni_para = self.doc.add_paragraph()
        uni_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        uni_run = uni_para.add_run("东北大学本科毕业论文")
        self._apply_run_font(uni_run, self.styles.font("cover_school_name"))

        self._add_cover_spacer(self.styles.layout.cover_spacer_title)

        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_para.paragraph_format.line_spacing = Pt(self.styles.layout.cover_line_spacing_pt)
        title_run = title_para.add_run(meta.title)
        self._apply_run_font(title_run, self.styles.font("cover_thesis_title"))

        self._add_cover_spacer(self.styles.layout.cover_spacer_info)

        info_rows = [(l, v) for l, v in [
            ("学 院 名 称", meta.college),
            ("专 业 名 称", meta.major),
            ("学 生 姓 名", meta.student_name),
            ("指 导 教 师", meta.advisor),
            ("副指导教师", meta.co_advisor),
        ] if v]
        if info_rows:
            self._add_cover_info_table(info_rows, separator="：")

        self._add_cover_spacer(self.styles.layout.cover_spacer_info)

        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date_run = date_para.add_run(f"{meta.year}年{meta.month}月")
        self._apply_run_font(date_run, self.styles.font("cover_date"))

    def _build_english_cover(self):
        meta = self.thesis.metadata

        self._add_cover_spacer(self.styles.layout.cover_spacer_english)

        eng_title_para = self.doc.add_paragraph()
        eng_title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        eng_title_para.paragraph_format.line_spacing = Pt(self.styles.layout.cover_line_spacing_pt)
        eng_title_run = eng_title_para.add_run(meta.english_title)
        self._apply_run_font(eng_title_run, self.styles.font("cover_english_title"))

        self._add_cover_spacer(self.styles.layout.cover_spacer_english_after_title)

        by_para = self.doc.add_paragraph()
        by_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        by_run = by_para.add_run(f"by {meta.student_name}")
        self._apply_run_font(by_run, self.styles.font("cover_english_detail"))

        self._add_cover_spacer(self.styles.layout.cover_spacer_english_after_title)

        table = self.doc.add_table(rows=2, cols=3)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        co_name = meta.co_advisor.split()[0] if meta.co_advisor else ""
        co_title = "Engineer" if "工程师" in (meta.co_advisor or "") else "Professor"
        advisor_title = "Associate Professor" if "副教授" in (meta.advisor or "") else "Professor"
        cells = [
            ("Supervisor:", advisor_title, meta.advisor.split()[0] if meta.advisor else ""),
            ("Associate Supervisor:", co_title, co_name),
        ]
        for row_idx, (label, title, name) in enumerate(cells):
            for col_idx, text in enumerate([label, title, name]):
                cell = table.cell(row_idx, col_idx)
                cell.text = ""
                self._clean_cell_paragraph(cell)
                run = cell.paragraphs[0].add_run(text)
                self._apply_run_font(run, self.styles.font("cover_english_detail"))

        self._add_cover_spacer(self.styles.layout.cover_spacer_info)

        uni_para = self.doc.add_paragraph()
        uni_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        uni_run = uni_para.add_run("Northeastern University")
        self._apply_run_font(uni_run, self.styles.font("cover_english_university"))

        month_name = _EN_MONTHS[int(meta.month)] if meta.month and meta.month.isdigit() else meta.month
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date_run = date_para.add_run(f"{month_name} {meta.year}")
        self._apply_run_font(date_run, self.styles.font("cover_english_date"))

    def _add_cover_spacer(self, pt: int):
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(pt)
        spacer.paragraph_format.space_after = Pt(0)

    def _remove_table_borders(self, table):
        tbl_element = table._tbl
        tblPr = tbl_element.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl_element.insert(0, tblPr)
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            borders.append(make_border(edge, "none", "0"))
        tblPr.append(borders)

    def _add_cover_info_table(self, rows, font_key="cover_info", separator="  "):
        table = self.doc.add_table(rows=len(rows), cols=2)
        self._remove_table_borders(table)

        font = self.styles.font(font_key)
        space_after = Pt(self.styles.layout.cover_detail_space_after_pt)

        for row_idx, (label, value) in enumerate(rows):
            lc = table.cell(row_idx, 0)
            vc = table.cell(row_idx, 1)

            self._clean_cell_paragraph(lc)
            lp = lc.paragraphs[0]
            lp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            lp.paragraph_format.space_after = Pt(0)
            lp.paragraph_format.space_before = Pt(0)
            lr = lp.add_run(f"{label}{separator}")
            self._apply_run_font(lr, font)

            self._clean_cell_paragraph(vc)
            vp = vc.paragraphs[0]
            vp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            vp.paragraph_format.space_after = space_after
            vp.paragraph_format.space_before = Pt(0)
            vr = vp.add_run(value)
            self._apply_run_font(vr, font)

    def _add_cover_detail_line(self, text: str, font_key: str = "cover_info"):
        para = self.doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(text)
        self._apply_run_font(run, self.styles.font(font_key))
        para.paragraph_format.space_after = Pt(self.styles.layout.cover_detail_space_after_pt)


    def _build_declaration(self):
        self._add_title_paragraph(self._T_DECLARATION, "declaration_title",
                                  style="Normal",
                                  space_before=self.styles.layout.declaration_title_space_lines,
                                  space_after=self.styles.layout.declaration_title_space_lines)

        decl_text = (
            "本人呈交的学位论文，是在导师的指导下，独立进行研究工作所取得的成果，"
            "所有数据、图片资料真实可靠。尽我所知，除文中已经注明引用的内容外，"
            "本学位论文的研究成果不包含他人享有著作权的内容。对本论文所涉及的"
            "研究工作做出贡献的其他个人和集体，均已在文中以明确的方式标明。"
            "本学位论文的知识产权归属于培养单位。"
        )
        decl_para = self.doc.add_paragraph()
        self._set_paragraph_text(decl_para, decl_text, "declaration_body", True)
        decl_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        decl_para.paragraph_format.first_line_indent = Pt(self.styles.layout.first_line_indent_pt)
        decl_para.paragraph_format.line_spacing = Pt(self.styles.layout.line_spacing_pt)
        decl_para.paragraph_format.space_after = Pt(0)
        decl_para.paragraph_format.space_before = Pt(0)

        sign_para = self.doc.add_paragraph()
        sign_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        self._set_paragraph_text(sign_para, "本人签名：________________      日期：_______________",
                                 "declaration_body", True)


    def _build_abstract_section(self, *, title, title_font, lines, body_font,
                                 indent=False, justify=False,
                                 kw_label, kw_label_font, keywords, kw_font, kw_sep):
        self._add_title_paragraph(title, title_font, style="Normal")

        for line in lines:
            if not line.strip():
                continue
            para = self.doc.add_paragraph()
            self._set_paragraph_text(para, line, body_font, True)
            if indent:
                para.paragraph_format.first_line_indent = Pt(self.styles.layout.first_line_indent_pt)
            if justify:
                para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            para.paragraph_format.space_after = Pt(self.styles.layout.abstract_body_space_after_pt)

        self.doc.add_paragraph()

        keyword_para = self.doc.add_paragraph()
        self._apply_line_spacing(keyword_para)
        title_run = keyword_para.add_run(kw_label)
        self._apply_run_font(title_run, self.styles.font(kw_label_font))
        if keywords:
            kw_run = keyword_para.add_run(kw_sep.join(keywords))
            self._apply_run_font(kw_run, self.styles.font(kw_font))


    def _build_toc(self):
        self._add_title_paragraph(self._T_TOC, "toc_title", style="Normal")

        self._add_toc_field()

    def _add_toc_field(self):
        para = self.doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(0)

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        run = para.add_run()
        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_separate)
        self._apply_run_font(run, self.styles.font("toc_other"))

        hint_run = para.add_run('目录（在Word中右键"更新域"刷新）')
        self._apply_run_font(hint_run, self.styles.font("toc_other"))
        hint_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        end_run = para.add_run()
        end_run._r.append(fld_end)
        self._apply_run_font(end_run, self.styles.font("toc_other"))


    def _add_seq_field(self, para, seq_name: str, reset: bool = False,
                        reset_val: int = 1, font_key: str = None):
        instr = f"SEQ {seq_name} \\* ARABIC"
        if reset:
            instr += f" \\r {reset_val}"

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instr
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        r1 = para.add_run()
        r1._r.append(fld_begin)
        r1._r.append(instr_text)
        r1._r.append(fld_sep)
        if font_key:
            self._apply_run_font(r1, self.styles.font(font_key))

        r2 = para.add_run(str(reset_val))
        r2._r.append(fld_end)
        if font_key:
            self._apply_run_font(r2, self.styles.font(font_key))

    _ASSIGNED_NUM_RE = re.compile(r'^([图表])(\d+)\.(\d+)$')

    def _add_field_caption(self, assigned_number: str, caption: str,
                            font_key: str, keep_with_next: bool = False):
        match = self._ASSIGNED_NUM_RE.match(assigned_number or '')
        if not match:
            text = caption_with_number(assigned_number, caption)
            self._add_caption_paragraph(text, font_key, keep_with_next)
            return

        prefix = match.group(1)
        chapter = match.group(2)
        sub_num = int(match.group(3))
        seq_name = "Figure" if prefix == "图" else "Table"

        para = self._make_caption_para(keep_with_next)
        native = self.styles.use_native_caption

        prefix_run = para.add_run(f"{prefix}{chapter}.")
        if not native:
            self._apply_run_font(prefix_run, self.styles.font(font_key))

        seq_font = font_key if not native else None
        self._add_seq_field(para, seq_name, reset=(sub_num == 1),
                            reset_val=sub_num, font_key=seq_font)

        if caption:
            cap_run = para.add_run(f" {caption}")
            if not native:
                self._apply_run_font(cap_run, self.styles.font(font_key))


    _HEADING_STYLES = frozenset({"Heading 1", "Heading 2", "Heading 3"})

    _LEVEL_CONFIG = {
        1: ("Heading 1", "chapter_title", WD_PARAGRAPH_ALIGNMENT.CENTER, "chapter_before_lines", "chapter_after_lines"),
        2: ("Heading 2", "section_title_1", WD_PARAGRAPH_ALIGNMENT.JUSTIFY, "section_before_lines", "section_after_lines"),
        3: ("Heading 3", "section_title_2", WD_PARAGRAPH_ALIGNMENT.JUSTIFY, "subsection_before_lines", "subsection_after_lines"),
    }

    def _add_heading(self, text: str, level: int):
        cfg = self._LEVEL_CONFIG[min(level, 3)]
        style_name, font_key, alignment, sb_key, sa_key = cfg
        para = self.doc.add_paragraph(text, style=style_name)
        para.alignment = alignment
        self._apply_line_spacing(para)
        set_spacing_lines(para, getattr(self.styles.layout, sb_key), getattr(self.styles.layout, sa_key))
        return para

    def _add_body_paragraph(self, text: str):
        para = self.doc.add_paragraph()
        self._add_body_with_citations(para, text)
        para.paragraph_format.line_spacing = Pt(self.styles.layout.line_spacing_pt)
        para.paragraph_format.first_line_indent = Pt(self.styles.layout.first_line_indent_pt)
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def _build_section(self, section: Section):
        if section.auto_number:
            heading_text = f"{section.auto_number} {section.title}"
        else:
            heading_text = section.title
        self._add_heading(heading_text, section.level)

        for item in section.items:
            if isinstance(item, str):
                if not item.strip():
                    continue
                self._add_body_paragraph(item)
            elif isinstance(item, Figure):
                self._add_figure(item)
            elif isinstance(item, Table):
                self._add_table(item)
            elif isinstance(item, PlantUMLBlock):
                self._add_plantuml(item)
            elif isinstance(item, CodeBlock):
                self._add_code(item)

        for subsection in section.subsections:
            self._build_section(subsection)

        if section.has_summary and section.summary_content:
            self._add_heading("本章小结", 2)
            for line in section.summary_content.split('\n'):
                if not line.strip():
                    continue
                self._add_body_paragraph(line.strip())


    def _build_references(self):
        self._add_title_paragraph(self._T_REFERENCES, "references_title")

        for ref in self.thesis.references:
            para = self.doc.add_paragraph()
            self._set_paragraph_text(para, ref.content, "references_body", True)

            char_pt = self.styles.font("references_body")["size"]
            pf = para.paragraph_format
            pf.left_indent = Pt(char_pt * 2)
            pf.first_line_indent = Pt(-char_pt * 2)
            pf.line_spacing = Pt(self.styles.layout.line_spacing_pt)
            pf.space_after = Pt(0)
            pf.space_before = Pt(0)

            bookmark_name = f"_Ref{ref.index}"
            bm_start = OxmlElement('w:bookmarkStart')
            bm_start.set(qn('w:id'), str(ref.index))
            bm_start.set(qn('w:name'), bookmark_name)
            bm_end = OxmlElement('w:bookmarkEnd')
            bm_end.set(qn('w:id'), str(ref.index))
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                pPr.addnext(bm_start)
            else:
                para._element.insert(0, bm_start)
            para._element.append(bm_end)


    def _build_appendix(self):
        self._add_title_paragraph(self._T_APPENDIX, "appendix_title")

        for section in self.thesis.appendix_sections:
            self._build_section(section)


    def _build_acknowledgments(self):
        self._add_title_paragraph(self._T_ACKNOWLEDGMENTS, "acknowledgments_title")

        for line in self.thesis.acknowledgments:
            if not line.strip():
                continue
            para = self.doc.add_paragraph()
            self._set_paragraph_text(para, line, "body", True)
            para.paragraph_format.first_line_indent = Pt(self.styles.layout.first_line_indent_pt)


    def _content_width_inches(self) -> float:
        return self._content_width_cm / 2.54

    def _content_height_inches(self) -> float:
        return self._content_height_cm / 2.54

    def _fit_image_size(self, img_source, scale: float = 1.0):
        """Compute (width, height) in Inches, fitting within content area * scale."""
        max_w = self._content_width_inches()
        max_h = self._content_height_inches()
        if PILImage is not None:
            try:
                img = PILImage.open(img_source)
                dpi = img.info.get("dpi", (96, 96))
                natural_w = img.size[0] / dpi[0]
                natural_h = img.size[1] / dpi[1]
                if hasattr(img_source, "seek"):
                    img_source.seek(0)
                w = min(natural_w, max_w)
                h = w * natural_h / natural_w
                if h > max_h:
                    h = max_h
                    w = h * natural_w / natural_h
                return Inches(w * scale), Inches(h * scale)
            except (OSError, ValueError):
                if hasattr(img_source, "seek"):
                    img_source.seek(0)
        else:
            print("Warning: PIL/Pillow not installed, images render without DPI-aware sizing or height clamping", file=sys.stderr)
        return Inches(max_w * scale), None

    def _add_image_with_caption(self, img_source, assigned_number: str,
                                caption: str, scale: float = 1.0,
                                missing_label: str = "图片"):
        img_para = self.doc.add_paragraph()
        img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        img_para.paragraph_format.keep_with_next = True
        effective_scale = scale if scale and scale > 0 else 1.0
        width, height = self._fit_image_size(img_source, effective_scale)
        if hasattr(img_source, "seek"):
            img_source.seek(0)
        try:
            img_para.add_run().add_picture(img_source, width=width, height=height)
        except Exception as exc:
            print(f"Warning: Failed to insert image '{missing_label}' ({exc})", file=sys.stderr)
            img_para._element.getparent().remove(img_para._element)
            label = getattr(img_source, 'name', str(img_source)) if not isinstance(img_source, str) else img_source
            self._add_error_paragraph(f"[{missing_label}: {label}]")
            return

        if assigned_number or caption:
            self._add_field_caption(assigned_number, caption, "figure_caption")

    def _add_figure(self, item):
        img_path = os.path.join(self.figures_dir, item.filename)
        self._add_image_with_caption(
            img_path, item.assigned_number,
            item.caption, item.scale, "图片",
        )

    def _add_table(self, item):
        if item.assigned_number or item.caption:
            self._add_field_caption(item.assigned_number, item.caption,
                                    "table_caption", keep_with_next=True)

        num_cols = len(item.headers)
        if num_cols == 0:
            return

        num_rows = 1 + len(item.rows)
        tbl = self.doc.add_table(rows=num_rows, cols=num_cols)
        tbl.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for row in tbl.rows:
            set_row_cant_split(row)

        set_row_table_header(tbl.rows[0])

        for col_idx, header in enumerate(item.headers):
            cell = tbl.cell(0, col_idx)
            cell.text = ""
            self._clean_cell_paragraph(cell)
            run = cell.paragraphs[0].add_run(header)
            self._apply_run_font(run, self.styles.font("table_caption"))
            run.font.bold = True

        for row_idx, row_data in enumerate(item.rows):
            for col_idx in range(num_cols):
                text = row_data[col_idx] if col_idx < len(row_data) else ""
                cell = tbl.cell(row_idx + 1, col_idx)
                cell.text = ""
                self._clean_cell_paragraph(cell)
                run = cell.paragraphs[0].add_run(text)
                self._apply_run_font(run, self.styles.font("table_caption"))

        self._apply_three_line_borders(tbl)

    def _apply_three_line_borders(self, tbl):
        tbl_element = tbl._tbl
        tblPr = tbl_element.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl_element.insert(0, tblPr)

        existing_borders = tblPr.find(qn("w:tblBorders"))
        if existing_borders is not None:
            tblPr.remove(existing_borders)

        borders = OxmlElement("w:tblBorders")
        borders.append(make_border("top", "single", "12"))
        borders.append(make_border("bottom", "single", "12"))
        for edge in ("left", "right", "insideV", "insideH"):
            borders.append(make_border(edge, "none", "0"))
        tblPr.append(borders)

        if len(tbl.rows) > 1:
            for cell in tbl.rows[1].cells:
                tc = cell._tc
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    tc.insert(0, tcPr)
                tcBorders = OxmlElement("w:tcBorders")
                tcBorders.append(make_border("top", "single", "6"))
                tcPr.append(tcBorders)

    @staticmethod
    def _clean_cell_paragraph(cell):
        for para in cell.paragraphs:
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                remove_child_tags(pPr, qn("w:widowControl"), qn("w:bidi"), qn("w:spacing"))
                rPr = pPr.find(qn("w:rPr"))
                if rPr is not None and len(rPr) == 0:
                    pPr.remove(rPr)

    def _add_code(self, code: CodeBlock):
        code_font = self.styles.font("code")
        para = self.doc.add_paragraph()
        self._apply_line_spacing(para)
        run = para.add_run(code.content)
        self._apply_run_font(run, code_font)

        name_para = self.doc.add_paragraph(f"// {code.filename}")
        self._apply_line_spacing(name_para)
        for run in name_para.runs:
            self._apply_run_font(run, {**code_font, "bold": False})
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    def _add_plantuml(self, item):
        uml_source = item.content.strip()
        if not uml_source:
            self._add_plantuml_error("PlantUML代码为空")
            return

        if "@startuml" not in uml_source:
            uml_source = f"@startuml\n{uml_source}\n@enduml"

        cmd = ["plantuml", "-charset", "UTF-8", "-tpng", "-pipe",
               "-DskinParam.backgroundColor=transparent"]
        try:
            result = subprocess.run(
                cmd, input=uml_source.encode("utf-8"),
                capture_output=True, check=False,
            )
        except FileNotFoundError:
            self._add_plantuml_error("未找到plantuml命令，请先安装PlantUML")
            return

        if result.returncode != 0:
            message = result.stderr.decode("utf-8", errors="ignore").strip() or "PlantUML渲染失败"
            self._add_plantuml_error(message)
            return

        if not result.stdout:
            self._add_plantuml_error("PlantUML未生成图像输出")
            return

        img_data = io.BytesIO(result.stdout)
        self._add_image_with_caption(
            img_data, item.assigned_number,
            item.caption, item.scale, "UML",
        )

    def _make_caption_para(self, keep_with_next: bool = False):
        if self.styles.use_native_caption:
            para = self.doc.add_paragraph(style="Caption")
        else:
            para = self.doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._apply_line_spacing(para)
        set_spacing_lines(para, self.styles.layout.caption_before_lines, self.styles.layout.caption_after_lines)
        if keep_with_next:
            para.paragraph_format.keep_with_next = True
        return para

    def _add_caption_paragraph(self, text: str, font_key: str = "figure_caption",
                               keep_with_next: bool = False):
        cap_para = self._make_caption_para(keep_with_next)
        cap_run = cap_para.add_run(text)
        if not self.styles.use_native_caption:
            self._apply_run_font(cap_run, self.styles.font(font_key))

    def _add_error_paragraph(self, message: str):
        para = self.doc.add_paragraph(message)
        if para.runs:
            para.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    def _add_plantuml_error(self, message: str):
        self._add_error_paragraph(f"[PlantUML渲染失败] {message}")


_CN_DIGITS = {'0': '○', '1': '一', '2': '二', '3': '三', '4': '四',
              '5': '五', '6': '六', '7': '七', '8': '八', '9': '九'}
_CN_MONTHS = {'1': '一', '2': '二', '3': '三', '4': '四', '5': '五', '6': '六',
              '7': '七', '8': '八', '9': '九', '10': '十', '11': '十一', '12': '十二'}
_EN_MONTHS = ["", "January", "February", "March", "April", "May", "June",
               "July", "August", "September", "October", "November", "December"]


def _chinese_year(year_str: str) -> str:
    return ''.join(_CN_DIGITS.get(d, d) for d in str(year_str))


def _chinese_month(month_str: str) -> str:
    return _CN_MONTHS.get(str(month_str), str(month_str))


